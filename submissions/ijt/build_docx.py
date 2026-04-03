"""Build IJT submission DOCX from manuscript.md.

Pipeline:
1. Pre-process: convert LaTeX math ($...$) to Unicode characters
2. pandoc: markdown → raw DOCX (with citeproc for references)
3. python-docx: post-process for APA formatting (TNR 12pt, double-spaced, margins)

Usage:
    .venv/bin/python submissions/ijt/build_docx.py
"""
import re
import subprocess
import sys
import tempfile
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_LINE_SPACING
from docx.oxml.ns import qn

SCRIPT_DIR = Path(__file__).parent
PROJECT_DIR = SCRIPT_DIR.parent.parent
MANUSCRIPT = SCRIPT_DIR / "manuscript.md"
REFERENCES = SCRIPT_DIR / "references.bib"
CSL = PROJECT_DIR.parent / "phd-thesis" / "csl" / "apa.csl"
OUTPUT = SCRIPT_DIR / "manuscript_with_author.docx"
OUTPUT_ANON = SCRIPT_DIR / "manuscript_anonymous.docx"
RAW_DOCX = SCRIPT_DIR / ".manuscript_raw.docx"

FONT_NAME = "Times New Roman"
BODY_SIZE = Pt(12)
TABLE_SIZE = Pt(10)

# --- Step 0: LaTeX math → Unicode ---

# Single-token LaTeX commands → Unicode
LATEX_TO_UNICODE = {
    r"\alpha": "α",
    r"\beta": "β",
    r"\gamma": "γ",
    r"\delta": "δ",
    r"\omega": "ω",
    r"\chi": "χ",
    r"\Phi": "Φ",
    r"\phi": "φ",
    r"\Delta": "Δ",
    r"\Sigma": "Σ",
    r"\sigma": "σ",
    r"\mu": "μ",
    r"\lambda": "λ",
    r"\kappa": "κ",
    r"\eta": "η",
    r"\epsilon": "ε",
    r"\tau": "τ",
    r"\rho": "ρ",
    r"\pi": "π",
    r"\leq": "≤",
    r"\geq": "≥",
    r"\approx": "≈",
    r"\pm": "±",
    r"\times": "×",
    r"\neq": "≠",
    r"\infty": "∞",
    r"\sim": "∼",
}

# Multi-token exact-match patterns (checked before single-token replacements)
LATEX_EXACT = {
    r"\chi^2": "χ²",
    r"\chi^{2}": "χ²",
    r"r_{ii}": "r\u1d62\u1d62",
    r"\Delta\beta": "Δβ",
}


def latex_to_unicode(text: str) -> str:
    """Replace LaTeX math expressions ($...$) with Unicode equivalents."""
    def replace_math(match):
        expr = match.group(1).strip()

        # Handle standalone minus sign
        if expr == "-":
            return "−"  # Unicode minus U+2212

        # Handle absolute value bars
        if expr == "|":
            return "|"

        # Try exact matches first
        if expr in LATEX_EXACT:
            return LATEX_EXACT[expr]

        # Try single-token replacements (longest first to avoid partial matches)
        result = expr
        for latex_cmd, unicode_char in sorted(
            LATEX_TO_UNICODE.items(), key=lambda x: -len(x[0])
        ):
            result = result.replace(latex_cmd, unicode_char)

        # Clean up remaining braces and backslashes
        result = result.replace("{", "").replace("}", "")
        result = result.replace("^2", "²")
        result = result.replace("^3", "³")

        return result

    # Match $...$ but not $$...$$
    return re.sub(r"(?<!\$)\$([^$]+)\$(?!\$)", replace_math, text)


def _preprocess_math(text: str) -> str:
    """Convert LaTeX math to Unicode and return processed text."""
    processed = latex_to_unicode(text)
    processed = processed.replace("$|Δ$", "|Δ")
    processed = re.sub(r"\|Δ([A-Z]+)\|", r"|Δ\1|", processed)
    return processed


def step0_preprocess() -> tuple[Path, Path]:
    """Pre-process markdown: math → Unicode, then create author + anonymous versions.

    Returns (author_path, anonymous_path).
    """
    text = MANUSCRIPT.read_text(encoding="utf-8")
    orig_math = len(re.findall(r"(?<!\$)\$([^$]+)\$(?!\$)", text))
    processed = _preprocess_math(text)
    remaining = len(re.findall(r"(?<!\$)\$([^$]+)\$(?!\$)", processed))
    print(f"  preprocess: {orig_math} math expressions → {remaining} remaining")

    # --- Author version ---
    author_path = SCRIPT_DIR / ".manuscript_author.md"
    author_path.write_text(processed, encoding="utf-8")

    # --- Anonymous version ---
    anon = processed

    # Remove author block (everything between YAML end and Abstract,
    # keeping running head, word count, and submission date)
    anon = re.sub(
        r"(---\n\n)(.*?)(## Abstract)",
        lambda m: m.group(1)
        + "\n".join(
            line
            for line in m.group(2).splitlines()
            if line.startswith("Running head:")
            or line.startswith("Word count:")
            or line.startswith("Submission date:")
            or line == ""
        )
        + "\n\n"
        + m.group(3),
        anon,
        flags=re.DOTALL,
    )

    # Remove author from YAML
    anon = re.sub(r'author: "Juda Kaleta"\n', "", anon)

    # Remove Biographical Note section entirely
    anon = re.sub(
        r"## Biographical Note\n\n.*?(?=\n## )",
        "",
        anon,
        flags=re.DOTALL,
    )

    # Anonymize OSF link
    anon = anon.replace("https://osf.io/yng37/", "[BLINDED]")
    anon = anon.replace("(OSF: [BLINDED])", "([BLINDED])")

    anon_path = SCRIPT_DIR / ".manuscript_anon.md"
    anon_path.write_text(anon, encoding="utf-8")

    return author_path, anon_path


# --- Step 1: pandoc conversion ---

def step1_pandoc(source: Path, output: Path):
    """Convert markdown to raw DOCX via pandoc."""
    cmd = [
        "pandoc", str(source),
        "-o", str(output),
        f"--bibliography={REFERENCES}",
        f"--csl={CSL}",
        "--citeproc",
    ]
    result = subprocess.run(cmd, capture_output=True, text=True)
    if result.returncode != 0:
        print(f"pandoc error: {result.stderr}", file=sys.stderr)
        sys.exit(1)
    print(f"  pandoc: {output.name} ({output.stat().st_size} bytes)")


# --- Step 2: python-docx post-processing ---

def step2_postprocess(raw: Path, final: Path):
    """Post-process DOCX with python-docx for APA formatting."""
    doc = Document(str(raw))

    # 1. Set margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # 2. Fix all styles
    for style in doc.styles:
        if hasattr(style, "font") and style.font is not None:
            style.font.name = FONT_NAME
            style.font.size = BODY_SIZE
        if hasattr(style, "paragraph_format") and style.paragraph_format is not None:
            style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
            style.paragraph_format.space_before = Pt(0)
            style.paragraph_format.space_after = Pt(0)

    # Heading styles
    for level in range(1, 6):
        name = f"Heading {level}"
        if name in doc.styles:
            s = doc.styles[name]
            s.font.name = FONT_NAME
            s.font.bold = True
            if level == 1:
                s.font.size = Pt(14)
                s.paragraph_format.space_before = Pt(12)
                s.paragraph_format.space_after = Pt(6)
            elif level == 2:
                s.font.size = Pt(13)
                s.paragraph_format.space_before = Pt(12)
                s.paragraph_format.space_after = Pt(6)
            else:
                s.font.size = BODY_SIZE
                s.paragraph_format.space_before = Pt(12)
                s.paragraph_format.space_after = Pt(6)

    # 3. Process all paragraphs
    for para in doc.paragraphs:
        for run in para.runs:
            run.font.name = FONT_NAME
            if run.font.size is None or run.font.size > Pt(16):
                pass  # Keep heading sizes
            else:
                run.font.size = BODY_SIZE
        para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE

    # 4. Process tables — APA table style
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                    para.paragraph_format.space_before = Pt(1)
                    para.paragraph_format.space_after = Pt(1)
                    for run in para.runs:
                        run.font.name = FONT_NAME
                        run.font.size = TABLE_SIZE

        tbl = table._tbl
        tblPr = tbl.find(qn("w:tblPr"))
        if tblPr is None:
            tblPr = tbl.makeelement(qn("w:tblPr"), {})
            tbl.insert(0, tblPr)

        # Set table width
        tblW = tblPr.find(qn("w:tblW"))
        if tblW is None:
            tblW = tblPr.makeelement(qn("w:tblW"), {})
            tblPr.append(tblW)
        tblW.set(qn("w:w"), str(int(Inches(6.5))))
        tblW.set(qn("w:type"), "dxa")

        # APA-style borders: top and bottom only
        tblBorders = tblPr.find(qn("w:tblBorders"))
        if tblBorders is not None:
            tblPr.remove(tblBorders)
        tblBorders = tblPr.makeelement(qn("w:tblBorders"), {})
        for border_name in ["top", "bottom"]:
            border = tblBorders.makeelement(
                qn(f"w:{border_name}"),
                {
                    qn("w:val"): "single",
                    qn("w:sz"): "4",
                    qn("w:space"): "0",
                    qn("w:color"): "000000",
                },
            )
            tblBorders.append(border)
        for border_name in ["left", "right", "insideV"]:
            border = tblBorders.makeelement(
                qn(f"w:{border_name}"),
                {qn("w:val"): "none", qn("w:sz"): "0", qn("w:space"): "0"},
            )
            tblBorders.append(border)
        insideH = tblBorders.makeelement(
            qn("w:insideH"),
            {qn("w:val"): "none", qn("w:sz"): "0", qn("w:space"): "0"},
        )
        tblBorders.append(insideH)
        tblPr.append(tblBorders)

        # Add bottom border to first row (header)
        if len(table.rows) > 0:
            for cell in table.rows[0].cells:
                tcPr = cell._tc.find(qn("w:tcPr"))
                if tcPr is None:
                    tcPr = cell._tc.makeelement(qn("w:tcPr"), {})
                    cell._tc.insert(0, tcPr)
                tcBorders = tcPr.makeelement(qn("w:tcBorders"), {})
                bottom = tcBorders.makeelement(
                    qn("w:bottom"),
                    {
                        qn("w:val"): "single",
                        qn("w:sz"): "4",
                        qn("w:space"): "0",
                        qn("w:color"): "000000",
                    },
                )
                tcBorders.append(bottom)
                tcPr.append(tcBorders)

    doc.save(str(final))
    raw.unlink()
    print(f"  output: {final.name} ({final.stat().st_size} bytes)")


# --- Main ---

if __name__ == "__main__":
    print("Building IJT manuscripts...")
    author_md, anon_md = step0_preprocess()

    raw_author = SCRIPT_DIR / ".raw_author.docx"
    raw_anon = SCRIPT_DIR / ".raw_anon.docx"

    print("\n  --- With author details ---")
    step1_pandoc(author_md, raw_author)
    step2_postprocess(raw_author, OUTPUT)

    print("\n  --- Anonymous ---")
    step1_pandoc(anon_md, raw_anon)
    step2_postprocess(raw_anon, OUTPUT_ANON)

    # Clean up temp markdown files
    author_md.unlink()
    anon_md.unlink()
    print("\nDone.")
